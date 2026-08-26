"""PDF / docx / txt / md extraction.

Routing logic:
  txt/md      → read directly
  docx        → python-docx
  pdf         → pdf-inspector classify:
                 text_based / mixed → local markdown (fast, free)
                 scanned / image_based → MinerU lightweight API (async)
"""
import logging
from pathlib import Path

import pdf_inspector

from . import mineru

logger = logging.getLogger(__name__)

SUPPORTED_EXT = {".txt", ".md", ".markdown", ".pdf", ".docx"}

# pdf-inspector PDF types that can be handled locally without OCR
_LOCAL_PDF_TYPES = {"text_based", "mixed"}


class UnsupportedTypeError(ValueError):
    pass


def parse_local(content: bytes, filename: str) -> str:
    """Synchronous local extraction. Returns raw text/markdown.

    Raises:
        UnsupportedTypeError: unsupported extension
        Exception: pdf-inspector / python-docx failure
    """
    ext = _ext(filename)
    if ext in {".txt", ".md", ".markdown"}:
        return content.decode("utf-8", errors="replace")
    if ext == ".docx":
        return _parse_docx(content)
    if ext == ".pdf":
        return _parse_pdf_local(content)
    raise UnsupportedTypeError(f"不支持的文件类型: {ext}")


def classify_pdf(content: bytes, filename: str) -> tuple[str, str | None]:
    """Classify a PDF and return (pdf_type, markdown_or_None).

    pdf_type: text_based | mixed | scanned | image_based | unknown
    If type is local-parseable, markdown is returned too (avoids double parse).
    """
    path = _temp_write(content, filename)
    try:
        try:
            result = pdf_inspector.process_pdf(path)
            pdf_type = str(result.pdf_type).lower()
            md = getattr(result, "markdown", None) or ""
        except Exception:  # pdf-inspector failed → treat as unknown, fall back to MinerU
            logger.exception("pdf-inspector failed for %s", filename)
            return "unknown", None
        if pdf_type in _LOCAL_PDF_TYPES:
            return pdf_type, md
        return pdf_type, None
    finally:
        import os

        os.unlink(path) if os.path.exists(path) else None


def _parse_pdf_local(content: bytes, filename: str = "doc.pdf") -> str:
    path = _temp_write(content, filename)
    try:
        result = pdf_inspector.process_pdf(path)
        return getattr(result, "markdown", "") or ""
    finally:
        import os

        if os.path.exists(path):
            os.unlink(path)


def _parse_docx(content: bytes) -> str:
    import docx
    import io

    doc = docx.Document(io.BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def _ext(filename: str) -> str:
    return Path(filename).suffix.lower()


def _temp_write(content: bytes, filename: str) -> str:
    import tempfile

    fd, path = tempfile.mkstemp(suffix=_ext(filename) or ".pdf")
    with open(fd, "wb") as f:
        f.write(content)
    return path
